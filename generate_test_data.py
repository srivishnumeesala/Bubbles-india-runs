# Technical Breakdown:
# - Generates synthetic test resumes and job descriptions for local end-to-end testing.
# - Creates Word (.docx) resumes containing explicit skills and proxy keywords (Harvard/Oxford).
# - Creates Job Description (.txt) file detailing core requirements.
# - Saves files directly into the configuration directories data/resumes and data/job_descriptions.

import os
from pathlib import Path
from docx import Document
from config.settings import RESUME_DIR, JD_DIR

def main():
    print("Generating synthetic test files...")
    
    # Ensure dirs exist
    RESUME_DIR.mkdir(parents=True, exist_ok=True)
    JD_DIR.mkdir(parents=True, exist_ok=True)

    # 1. Generate Priya Sharma Resume (High Fit, Prestige School)
    doc_priya = Document()
    doc_priya.add_heading("Priya Sharma", 0)
    doc_priya.add_paragraph("Summary: Experienced Python developer specializing in machine learning, PyTorch, and SQL.")
    doc_priya.add_paragraph("Education: studied at Harvard University. Graduated in 2014.")
    
    # Table for skills
    t_priya = doc_priya.add_table(rows=3, cols=2)
    t_priya.cell(0, 0).text = "Programming"
    t_priya.cell(0, 1).text = "Python, SQL, C++"
    t_priya.cell(1, 0).text = "Libraries"
    t_priya.cell(1, 1).text = "PyTorch, TensorFlow, Pandas"
    t_priya.cell(2, 0).text = "Frameworks"
    t_priya.cell(2, 1).text = "FastAPI, Git, Docker"
    
    doc_priya.save(RESUME_DIR / "priya_sharma.docx")
    print(f"Created: {RESUME_DIR / 'priya_sharma.docx'}")

    # 2. Generate Alex Johnson Resume (Mid Fit, Prestige School)
    doc_alex = Document()
    doc_alex.add_heading("Alex Johnson", 0)
    doc_alex.add_paragraph("Summary: Software engineer with experience building web applications and REST APIs using Python.")
    doc_alex.add_paragraph("Education: Studied at Oxford University. Graduated in 2012.")
    
    # Skills list
    doc_alex.add_paragraph("Skills:")
    doc_alex.add_paragraph("- Languages: Python, JavaScript, SQL")
    doc_alex.add_paragraph("- Frameworks: FastAPI, React")
    doc_alex.add_paragraph("- Tools: Git, Docker")
    
    doc_alex.save(RESUME_DIR / "alex_johnson.docx")
    print(f"Created: {RESUME_DIR / 'alex_johnson.docx'}")

    # 3. Generate Job Description text file
    jd_content = (
        "Job Description: Senior Python Developer\n\n"
        "We are looking for a software engineer to build data pipelines and design REST APIs.\n"
        "Core skills required:\n"
        "- Python\n"
        "- PyTorch\n"
        "- SQL\n"
        "- FastAPI\n"
        "- Docker\n\n"
        "Qualifications:\n"
        "- Bachelor's degree in Computer Science.\n"
        "- Good communication and Agile collaboration skills."
    )
    
    jd_path = JD_DIR / "senior_python_developer.txt"
    with open(jd_path, "w", encoding="utf-8") as f:
        f.write(jd_content)
    print(f"Created: {jd_path}")
    print("Done!")

if __name__ == "__main__":
    main()
