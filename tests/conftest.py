# Technical Breakdown:
# - Shared test fixtures for Pytest.
# - Generates valid, temporary PDF files using PyMuPDF (fitz) to avoid external assets.
# - Generates valid, temporary DOCX files using python-docx.
# - Provides clean mockup data for JDs and candidate resumes.

import pytest
from pathlib import Path
import fitz
from docx import Document

@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Generate a temporary PDF file with mock content."""
    pdf_path = tmp_path / "resume.pdf"
    doc = fitz.open()
    
    # Page 1
    page1 = doc.new_page()
    page1.insert_text((50, 50), "John Doe Resume\nEmail: john.doe@example.com\nWeb: www.johndoe.com\nHeader Line")
    page1.insert_text((50, 150), "Skills: Python, TensorFlow, SQL")
    
    # Page 2
    page2 = doc.new_page()
    page2.insert_text((50, 50), "Work Experience\nSenior Engineer at Tech Corp\nHeader Line")
    
    # Page 3
    page3 = doc.new_page()
    page3.insert_text((50, 50), "Education\nB.S. Computer Science\nHeader Line")
    
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path

@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Generate a temporary DOCX file with mock content including a table."""
    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Jane Smith Resume")
    doc.add_paragraph("Skills:")
    
    # Add a table (often used for layout in resumes)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Language"
    table.cell(0, 1).text = "Python, SQL"
    table.cell(1, 0).text = "Frameworks"
    table.cell(1, 1).text = "PyTorch, FastAPI"
    
    doc.save(str(docx_path))
    return docx_path

@pytest.fixture
def sample_jd() -> str:
    """Return a mock job description."""
    return """
    Job Description: Senior Python Developer
    We are looking for a software engineer with expertise in Python, SQL, and PyTorch.
    You will build data pipelines and design REST APIs using FastAPI or Flask.
    Qualifications: Bachelor's degree in Computer Science. AWS certification is a plus.
    """
